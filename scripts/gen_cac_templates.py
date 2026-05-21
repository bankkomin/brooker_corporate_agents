"""Generate CFO-friendly CAC templates: an Excel data pack + a Word monthly report.

The .md templates in obsidian-vault/cac/_templates are for the agent/wiki; the
CFO works in Excel/Word. This script renders the same mandate coverage into:
  - config/templates/cac/CAC_Monthly_Data_Pack.xlsx  (fill-in intake)
  - config/templates/cac/CAC_Monthly_CFO_Report.docx (output report/email body)
"""
from __future__ import annotations

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

XLSX = "config/templates/cac/CAC_Monthly_Data_Pack.xlsx"
DOCX = "config/templates/cac/CAC_Monthly_CFO_Report.docx"

H1 = Font(bold=True, size=14, color="FFFFFF")
HEAD = Font(bold=True, color="FFFFFF")
BOLD = Font(bold=True)
ITAL = Font(italic=True)
hdr_fill = PatternFill("solid", fgColor="1F4E78")
sub_fill = PatternFill("solid", fgColor="2E75B6")
input_fill = PatternFill("solid", fgColor="FFF2CC")   # yellow = fill me
ref_fill = PatternFill("solid", fgColor="E2EFDA")     # green = reference
thin = Side(style="thin", color="BFBFBF")
border = Border(left=thin, right=thin, top=thin, bottom=thin)
wrap = Alignment(wrap_text=True, vertical="top")
center = Alignment(horizontal="center", vertical="center")


def title_block(ws, title):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
    c = ws.cell(row=1, column=1, value=title)
    c.font = H1
    c.fill = hdr_fill
    c.alignment = Alignment(vertical="center", horizontal="left", indent=1)
    ws.row_dimensions[1].height = 24


def put_table(ws, start_row, headers, rows, input_cols=None, ref_cols=None):
    input_cols = input_cols or set()
    ref_cols = ref_cols or set()
    for j, h in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=j, value=h)
        cell.font = HEAD
        cell.fill = sub_fill
        cell.border = border
        cell.alignment = center
    r = start_row + 1
    for row in rows:
        for j, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=j, value=val)
            cell.border = border
            cell.alignment = wrap
            if j in input_cols and (val in (None, "")):
                cell.fill = input_fill
            elif j in ref_cols:
                cell.fill = ref_fill
        r += 1
    return r


def build_xlsx():
    wb = openpyxl.Workbook()

    ws = wb.active
    ws.title = "Instructions"
    title_block(ws, "CAC Monthly Data Pack  -  Month: __________")
    lines = [
        "",
        "Mandate: Manage the Group's balance sheet, liquidity, funding structure, capital",
        "allocation priorities, and asset-liability risk exposures in accordance with",
        "Board-approved strategy, risk appetite, and treasury policies.",
        "",
        "HOW TO USE",
        "  - Fill every YELLOW cell with the month-end figure. Leave blank only if N/A.",
        "  - GREEN cells are Board-approved targets/limits - do NOT edit them.",
        "  - Return to the CAC by the 5th business day after month-end close.",
        "  - Prepared by: __________   Date: __________",
        "",
        "GOVERNING BASIS",
        "  - Board strategy: Three-Engine Model + North Star 2028 (recurring THB 500M, AUM USD 600M, D/E <0.5x).",
        "  - Risk appetite: investment ratio <40% (by 30 Jun 2026), D/E <0.5x, cost of funds <3-4%,",
        "    on-chain DTV <=25% (hard cap; new loans LTV <10%), >=6 months offline-cash runway,",
        "    Sovereignty Buffer >= 100 BTC, treasury native yield >=10%.",
        "  - Treasury policies: Stay Liquid doctrine, melt-up/sell-down plan, Decision Rights Matrix (R-05).",
        "",
        "SHEETS: 1 Balance Sheet | 2 Capital Allocation | 3 Liquidity | 4 Funding Structure |",
        "        5 Asset-Liability | 6 On-Chain Collateral | 7 Risk Appetite | 8 Policy Compliance | 9 Recommendations",
    ]
    for i, t in enumerate(lines, 3):
        cell = ws.cell(row=i, column=1, value=t)
        if t.strip() and t.isupper():
            cell.font = BOLD
    ws.column_dimensions["A"].width = 100

    ws = wb.create_sheet("1 Balance Sheet")
    title_block(ws, "1. Balance Sheet  (mandate: balance sheet)")
    r = put_table(ws, 3,
        ["Line", "This month (THB M)", "Prior month", "MoM change"],
        [["Total assets", "", "", ""],
         ["  Current / liquid assets", "", "", ""],
         ["  Investments (principal + treasury)", "", "", ""],
         ["  Operating / other assets", "", "", ""],
         ["Total liabilities", "", "", ""],
         ["  Bank debt / facilities drawn", "", "", ""],
         ["  On-chain / structured borrowing", "", "", ""],
         ["  Other liabilities", "", "", ""],
         ["Total equity / Net Worth", "", "", ""]],
        input_cols={2, 3, 4})
    r += 1
    put_table(ws, r,
        ["Key ratio", "Value", "Target", "Status"],
        [["Debt-to-Equity (D/E)", "", "<0.5x", ""],
         ["Net Worth (THB M)", "", "grow", ""],
         ["Investment / Total Assets", "", "<40% (Jun 2026)", ""]],
        input_cols={2, 4}, ref_cols={3})
    ws.column_dimensions["A"].width = 38
    for col in "BCD":
        ws.column_dimensions[col].width = 18

    ws = wb.create_sheet("2 Capital Allocation")
    title_block(ws, "2. Capital Allocation Priorities")
    put_table(ws, 3,
        ["Bucket", "Actual (THB M)", "% of total assets", "Target range", "Variance"],
        [["Core operating businesses", "", "", "per strategy", ""],
         ["Principal investments", "", "", "<=10% non-core", ""],
         ["VCC Platform (Engine 1, GP/seed)", "", "", "~15%", ""],
         ["Advisory book (Engine 2)", "", "", "10-30%", ""],
         ["Digital Asset Treasury (Engine 3)", "", "", "20-30%", ""],
         ["Strategic reserves / cash", "", "", ">=6mo burn", ""],
         ["TOTAL", "", "", "100%", ""]],
        input_cols={2, 3, 5}, ref_cols={4})
    ws.column_dimensions["A"].width = 34
    for col in "BCDE":
        ws.column_dimensions[col].width = 16

    ws = wb.create_sheet("3 Liquidity")
    title_block(ws, "3. Liquidity & 'Stay Liquid' Doctrine")
    r = put_table(ws, 3,
        ["Item", "Value", "Threshold"],
        [["Operating cash (THB M)", "", "-"],
         ["Near-cash / liquid securities (THB M)", "", "-"],
         ["Sovereignty Buffer (BTC)", "", "floor: 100 BTC"],
         ["Monthly operating burn (THB M)", "", "budget ~8"],
         ["Months of runway", "", "min 6 months"]],
        input_cols={2}, ref_cols={3})
    r += 1
    ws.cell(row=r, column=1, value="Stress scenarios").font = BOLD
    r += 1
    r = put_table(ws, r,
        ["Scenario", "Assumption", "Result", "Pass/Fail"],
        [["Crypto drawdown", "-25% DAT", "", ""],
         ["Cash drain", "-50% cash MoM", "", ""],
         ["Combined", "-25% DAT + facility pull", "", ""]],
        input_cols={3, 4}, ref_cols={2})
    r += 1
    ws.cell(row=r, column=1, value="Contingency Funding Plan status:").font = BOLD
    ws.cell(row=r, column=2).fill = input_fill
    ws.column_dimensions["A"].width = 38
    for col in "BCD":
        ws.column_dimensions[col].width = 20

    ws = wb.create_sheet("4 Funding Structure")
    title_block(ws, "4. Funding Structure & Facilities")
    hdrs = ["Lender", "Type", "Limit (THB M)", "Drawn", "Available", "Util %",
            "All-in rate", "Maturity", "Secured?", "Covenant", "Headroom", "Renewal status"]
    rows = [["SCB", "", "", "", "", "", "", "", "", "(lender covenant)", "", ""],
            ["BBL", "", "", "", "", "", "", "", "", "", "", ""],
            ["(other)", "", "", "", "", "", "", "", "", "", "", ""],
            ["AGGREGATE", "-", "", "", "", "", "wtd", "-", "-", "-", "-", "-"]]
    put_table(ws, 3, hdrs, rows, input_cols=set(range(2, 13)))
    r = 8
    ws.cell(row=r, column=1, value="Funding mix:").font = BOLD
    ws.cell(row=r, column=2, value="bank __% / on-chain __% / bond __% / equity __%")
    ws.cell(row=r + 1, column=1, value="Covenant watch:").font = BOLD
    ws.cell(row=r + 1, column=2, value="any covenant within 10% of limit? maturity <90 days w/o renewal?")
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["J"].width = 22
    for col in "BCDEFGHIKL":
        ws.column_dimensions[col].width = 13

    ws = wb.create_sheet("5 Asset-Liability")
    title_block(ws, "5. Asset-Liability Risk Exposures")
    r = put_table(ws, 3,
        ["Tenor bucket", "Assets (THB M)", "Liabilities (THB M)", "Gap"],
        [["Overnight-1m", "", "", ""], ["1-3m", "", "", ""], ["3-12m", "", "", ""],
         ["1-5y", "", "", ""], [">5y", "", "", ""]],
        input_cols={2, 3, 4})
    r += 1
    for label, ref in [("Weighted asset duration (yr):", ""),
                       ("Weighted liability duration (yr):", ""),
                       ("Duration gap (yr):", "target <2.0"),
                       ("Refinancing due next 12m (THB M + plan):", ""),
                       ("Collateral sufficiency on secured assets (%):", ""),
                       ("FX exposure (USD treasury vs THB liab):", "")]:
        ws.cell(row=r, column=1, value=label).font = BOLD
        ws.cell(row=r, column=2).fill = input_fill
        if ref:
            c = ws.cell(row=r, column=3, value=ref)
            c.fill = ref_fill
        r += 1
    ws.column_dimensions["A"].width = 40
    for col in "BCD":
        ws.column_dimensions[col].width = 18

    ws = wb.create_sheet("6 On-Chain Collateral")
    title_block(ws, "6. On-Chain Collateral & Margin")
    put_table(ws, 3,
        ["Protocol/Venue", "Collateral asset", "Collateral (USD)", "Borrowed (USD)",
         "DTV %", "Liquidation price", "Buffer to liq", "Tier (<10/10-25/>25%)"],
        [["", "", "", "", "", "", "", ""], ["", "", "", "", "", "", "", ""],
         ["AGGREGATE DTV", "-", "", "", "", "-", "-", "-"]],
        input_cols={1, 2, 3, 4, 5, 6, 7, 8})
    ws.cell(row=7, column=1,
            value="DTV tiers per R-05: <10% management | 10-25% committee | >25% Board.").font = ITAL
    for i, w in enumerate([18, 16, 16, 16, 10, 16, 14, 20], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws = wb.create_sheet("7 Risk Appetite")
    title_block(ws, "7. Risk Appetite - Limits & Escalation")
    put_table(ws, 3,
        ["Metric", "Current", "Limit (appetite)", "Status", "Escalate to"],
        [["Investment / Total Assets", "", "40%", "", "Board (post-Jun 2026)"],
         ["Debt-to-Equity", "", "<0.5x", "", "Board"],
         ["Cost of funds", "", "<3-4%", "", "CAC"],
         ["Aggregate on-chain DTV", "", "<=25% (hard cap)", "", "Board if >25%"],
         ["New on-chain loan LTV", "", "<10%", "", "Committee if >=10%"],
         ["Treasury native yield", "", ">=10% / THB 100M", "", "CAC"],
         ["Liquidity runway", "", ">=6mo offline cash", "", "CEO if <6mo"],
         ["Sovereignty Buffer (BTC)", "", ">=100", "", "CEO if breached"]],
        input_cols={2, 4}, ref_cols={3, 5})
    ws.column_dimensions["A"].width = 26
    for col in "BCDE":
        ws.column_dimensions[col].width = 18

    ws = wb.create_sheet("8 Policy Compliance")
    title_block(ws, "8. Policy Compliance Confirmation")
    put_table(ws, 3,
        ["Policy / strategy", "Compliant? (Y/N)", "Note"],
        [["Board-approved strategy (Three-Engine / North Star)", "", ""],
         ["Stay Liquid treasury doctrine", "", ""],
         ["Melt-up / sell-down plan (40% glide path)", "", ""],
         ["Decision Rights Matrix (R-05) respected", "", ""]],
        input_cols={2, 3})
    ws.column_dimensions["A"].width = 48
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 40

    ws = wb.create_sheet("9 Recommendations")
    title_block(ws, "9. Material Changes Recommended This Month")
    put_table(ws, 3,
        ["Recommendation", "Amount", "Required approval (R-05)", "Rationale"],
        [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
        input_cols={1, 2, 3, 4})
    ws.column_dimensions["A"].width = 34
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 40

    wb.save(XLSX)
    print("SAVED", XLSX, "sheets:", len(wb.sheetnames))


def build_docx():
    import docx
    from docx.shared import Pt, RGBColor

    d = docx.Document()
    d.add_heading("CAC Monthly Report", level=0)
    p = d.add_paragraph()
    p.add_run("To: ").bold = True
    p.add_run("Supane, CFO     ")
    p.add_run("From: ").bold = True
    p.add_run("Capital Allocation & ALCO Committee (CAC)")
    p = d.add_paragraph()
    p.add_run("Period: ").bold = True
    p.add_run("__________     ")
    p.add_run("Prepared: ").bold = True
    p.add_run("__________")
    d.add_paragraph(
        "Mandate: manage the Group's balance sheet, liquidity, funding structure, capital "
        "allocation priorities, and asset-liability risk exposures in accordance with "
        "Board-approved strategy, risk appetite, and treasury policies."
    ).italic = True

    def H(t):
        d.add_heading(t, level=1)

    def fill(t):
        para = d.add_paragraph(t)
        return para

    H("1. Executive Summary")
    fill("[3-5 sentences: overall balance-sheet health, the single biggest change this "
         "month, and any limit breaches/approaches needing CFO or Board attention.]")

    H("2. Balance Sheet")
    fill("Total assets: THB ___M | Total liabilities: THB ___M | Net Worth: THB ___M")
    fill("Debt-to-Equity: ___x (target <0.5x) - [status]")
    fill("MoM movement & drivers: [____]")

    H("3. Capital Allocation")
    fill("Investment-to-total-assets ratio: ___% (ceiling 40% by 30 Jun 2026; OKR <38%).")
    fill("Allocation vs Three-Engine targets: [one line per engine].")
    fill("Material drift from target: [____]")

    H("4. Liquidity - 'Stay Liquid'")
    fill("Runway: ___ months (min 6) | Sovereignty Buffer: ___ BTC (floor 100).")
    fill("Stress-test outcome: [pass/fail]. Contingency Funding Plan: [status].")

    H("5. Funding & Covenants")
    fill("Total drawn / available: THB ___ / ___M (util ___%).")
    fill("Leverage: D/E ___x vs <0.5x; cost of funds ___% vs <3-4%. Maturities <90 days / renewal: [____].")
    fill("Bank relationship notes (SCB/BBL): [____].")

    H("6. Asset-Liability Risk")
    fill("Duration gap: ___ yr (target <2.0). Refinancing due next 12m: THB ___M.")
    fill("Collateral coverage: ___%. On-chain aggregate DTV: ___% (Board review >25%).")

    H("7. Limit Breaches & Escalations")
    tbl = d.add_table(rows=2, cols=4)
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(["Metric", "Value", "Limit", "Action"]):
        tbl.rows[0].cells[j].text = h
    for j in range(4):
        tbl.rows[1].cells[j].text = ""
    d.add_paragraph("(If none: 'No limits breached this month.')").italic = True

    H("8. Policy Compliance")
    fill("Confirm compliant (Y/N) with: Board strategy (Three-Engine/North Star); Stay "
         "Liquid doctrine; melt-up/sell-down 40% glide path; Decision Rights Matrix (R-05).")

    H("9. Recommendations to CFO / IC / Board")
    tbl = d.add_table(rows=2, cols=3)
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(["Recommendation", "Amount", "Required approval (R-05)"]):
        tbl.rows[0].cells[j].text = h
    for j in range(3):
        tbl.rows[1].cells[j].text = ""

    d.add_paragraph()
    foot = d.add_paragraph("Detail & source figures: see this month's CAC Monthly Data Pack "
                           "(Excel). Questions: #cac-committee or the CAC agent.")
    foot.italic = True

    d.save(DOCX)
    print("SAVED", DOCX)


if __name__ == "__main__":
    build_xlsx()
    build_docx()
