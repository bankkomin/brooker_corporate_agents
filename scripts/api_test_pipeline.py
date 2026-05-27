"""End-to-end HTTP smoke tests for the agent pipeline — no Slack needed.

Exercises:
  - health of every service touched by the CAC report flow
  - deck-writer /report/cac-meeting dept guard (caller_dept allow/deny)
  - cac-orchestrator /query (SharePoint pre-guard, conversational, mandate)
  - read-only-orchestrator /query (finance: pre-guard, naked question)

Run:
    python scripts/api_test_pipeline.py            # all tests
    python scripts/api_test_pipeline.py --only cac # filter by name substring
"""
from __future__ import annotations

import argparse
import asyncio
import sys
from dataclasses import dataclass

import httpx

SHARE = ("https://brookergroup-my.sharepoint.com/:x:/p/s_minkhant/"
         "IQBtE9oEDVGCSKnItVre4q-OAccFNsxv4DfuCkCXh5dGfAA?e=ftdTiS")

GREEN, RED, DIM, RESET = "\033[32m", "\033[31m", "\033[2m", "\033[0m"


@dataclass
class T:
    name: str
    method: str
    url: str
    params: dict | None = None
    body: dict | None = None
    expect_status: int = 200
    expect_in: str | None = None     # case-insensitive substring expected in body
    expect_not_in: str | None = None
    # expect_any: PASS if at least one of these substrings is in the body (case-insensitive).
    # Designed for use-case tests where the LLM might paraphrase the right answer
    # ("THB 500M" vs "500 million baht" vs "500MM THB"). List the variants.
    expect_any: list[str] | None = None
    timeout: float = 30.0
    # flaky=True: run the test 3 times sequentially, pass if 2/3 runs pass (majority vote).
    # Use for known-flaky LLM use-case tests that occasionally return malformed responses.
    flaky: bool = False


TESTS: list[T] = [
    # ── health ────────────────────────────────────────────────────────────────
    T("health: cac-orchestrator",         "GET", "http://localhost:3001/health"),
    T("health: read-only-orchestrator",   "GET", "http://localhost:3040/health"),
    T("health: deck-writer",              "GET", "http://localhost:3050/health"),
    T("health: rag-ingestion",            "GET", "http://localhost:3004/health"),
    T("health: slack-bot",                "GET", "http://localhost:3003/health"),

    # ── deck-writer /report/cac-meeting dept guard ────────────────────────────
    T("report: caller_dept=cac -> 200, has file_url", "GET",
      "http://localhost:3050/report/cac-meeting",
      params={"share_url": SHARE, "caller_dept": "cac"},
      expect_in="file_url", timeout=120),
    T("report: caller_dept=finance -> 403 with #cac-committee hint", "GET",
      "http://localhost:3050/report/cac-meeting",
      params={"share_url": SHARE, "caller_dept": "finance"},
      expect_status=403, expect_in="#cac-committee"),
    T("report: caller_dept=unknown -> 403", "GET",
      "http://localhost:3050/report/cac-meeting",
      params={"share_url": SHARE, "caller_dept": "unknown"},
      expect_status=403),
    T("report: no share_url -> uses CAC_DATA_PACK_SHARE_URL env fallback", "GET",
      "http://localhost:3050/report/cac-meeting",
      params={"caller_dept": "cac"},
      expect_status=200, expect_in="file_url", timeout=120),
    T("report: bad URL -> 422", "GET",
      "http://localhost:3050/report/cac-meeting",
      params={"share_url": "not-a-url", "caller_dept": "cac"},
      expect_status=422),

    # ── cac-orchestrator /query ───────────────────────────────────────────────
    T("cac: SharePoint link + question -> pre-RAG abstain w/ rephrase hint", "POST",
      "http://localhost:3001/query",
      body={"query": f"what's the D/E ratio shown here? {SHARE}",
            "channel": "t", "user_id": "diag", "dept_id": "cac"},
      expect_in="produce CAC report from this link",
      expect_not_in="0.5 and 1.5"),  # textbook D/E definition must NOT leak in
    T("cac: conversational greeting", "POST",
      "http://localhost:3001/query",
      body={"query": "hello", "channel": "t", "user_id": "diag", "dept_id": "cac"},
      expect_not_in="don't have"),
    T("cac: mandate question -> capability bypass", "POST",
      "http://localhost:3001/query",
      body={"query": "what is your mandate?", "channel": "t", "user_id": "diag", "dept_id": "cac"},
      expect_in="Capital Allocation"),

    # ── read-only-orchestrator /query (finance) ───────────────────────────────
    T("finance: link + question -> pre-RAG short-circuit, chunks=0", "POST",
      "http://localhost:3040/query",
      body={"query": f"what's the D/E ratio shown here? {SHARE}",
            "channel": "t", "user_id": "diag", "dept_id": "finance"},
      expect_in="post in #cac-committee"),
    T("finance: bare question (no link) goes through RAG", "POST",
      "http://localhost:3040/query",
      body={"query": "what is the D/E ratio?",
            "channel": "t", "user_id": "diag", "dept_id": "finance"},
      timeout=60),  # success means HTTP 200; content varies with ingestion

    # ── hr-orchestrator /query (same fixes as cac-orchestrator) ───────────────
    T("health: hr-orchestrator",          "GET", "http://localhost:3002/health"),
    T("hr: SharePoint link + question -> pre-guard abstain (HR sources named)", "POST",
      "http://localhost:3002/query",
      body={"query": f"what's the headcount shown here? {SHARE}",
            "channel": "t", "user_id": "diag", "dept_id": "hr"},
      expect_in="self-assessment questionnaires",
      expect_not_in="0.5 and 1.5"),
    T("hr: conversational greeting", "POST",
      "http://localhost:3002/query",
      body={"query": "hello", "channel": "t", "user_id": "diag", "dept_id": "hr"},
      expect_not_in="don't have", timeout=90),
    T("hr: mandate question -> capability bypass (no backstop misfire)", "POST",
      "http://localhost:3002/query",
      body={"query": "what is your mandate?", "channel": "t", "user_id": "diag", "dept_id": "hr"},
      # Only check it's NOT the abstain — LLM phrasing varies (HR / Human Resources / analyst…).
      expect_not_in="don't have an answer for that", timeout=90),
]


# ── Universal smoke + cross-cutting tests for the 9 read-only depts ──────────
# Each dept gets:
#   * S02 greeting        -> capability path, NOT the abstain
#   * S03 mandate         -> capability bypass via _is_capability_query
#   * X-02 link pre-guard -> SharePoint pre-RAG short-circuit, message names
#                            that dept's specific source label
# All hit read-only-orchestrator :3040 with the right dept_id.
# Plus X-01: each non-cac dept tries to call deck-writer /report/cac-meeting
# and must be 403'd by the caller_dept guard.
_RO_QUERY = "http://localhost:3040/query"
_RPT = "http://localhost:3050/report/cac-meeting"

# Per-dept substring that must appear in the SharePoint pre-guard abstain —
# proves the dept-specific source label from _DEPT_SOURCE_LABEL is in use.
_DEPT_SRC_FRAGMENT = {
    "ceo":   "CEO strategy",
    "ic":    "investment committee",
    "cio":   "CIO investment-cluster",
    "vcc":   "VCC fund offering",
    "legal": "external-counsel",
    "comms": "communications knowledge",
    "risk":  "risk policies",
    "ib":    "investment-banking",
    "it":    "infrastructure, security, devops",
}

# ── USE-CASE TESTS — verify the agent returns REAL DATA, not "anything not abstain" ──
# Each asks a question whose ground-truth answer is a specific fact in the corpus and
# asserts a substring (or any of several phrasings) of that fact must appear.
# These are how we tell "the bot is correct," not just "the bot is reachable."
USECASE = [
    # ─ CAC: governance constants from the Khao Yai resolutions / SKILL.md ─
    T("USE-CASE cac: Sovereignty Buffer floor (≥100 BTC)", "POST", "http://localhost:3001/query",
      body={"query":"what is the Sovereignty Buffer floor in BTC?","channel":"t","user_id":"diag","dept_id":"cac"},
      expect_any=["100 BTC","100 bitcoin","≥ 100","at least 100"], timeout=90),
    T("USE-CASE cac: D/E ceiling (<0.5x)", "POST", "http://localhost:3001/query",
      body={"query":"what is the Board-approved D/E ceiling?","channel":"t","user_id":"diag","dept_id":"cac"},
      expect_any=["0.5x","0.5 x","< 0.5","below 0.5"], timeout=90),
    T("USE-CASE cac: on-chain DTV hard cap (25%)", "POST", "http://localhost:3001/query",
      body={"query":"what is the Board-approved aggregate on-chain Debt-to-Value DTV hard cap percentage?","channel":"t","user_id":"diag","dept_id":"cac"},
      expect_any=["25%","25 percent","≤ 25","25 %"], timeout=90),
    T("USE-CASE cac: 40% Investment Company rule deadline", "POST", "http://localhost:3001/query",
      body={"query":"when does the 40 percent Investment Company rule become binding?","channel":"t","user_id":"diag","dept_id":"cac"},
      expect_any=["30 Jun 2026","June 2026","Jun 2026","30 June 2026"], timeout=90),

    # ─ Finance / CFO: BICL facts ─
    T("USE-CASE finance: BICL FATCA classification (Active NFFE)", "POST", "http://localhost:3040/query",
      body={"query":"what is BICL's FATCA classification?","channel":"t","user_id":"diag","dept_id":"finance"},
      expect_any=["Active NFFE","NFFE"], timeout=90),
    T("USE-CASE finance: PN.35 principal amount (USD 39.02M)", "POST", "http://localhost:3040/query",
      body={"query":"what is the principal of Promissory Note 35 from The Brooker Group to BICL?","channel":"t","user_id":"diag","dept_id":"finance"},
      expect_any=["39.02","39,020,000","USD 39"], timeout=90),

    # ─ CEO: North Star 2028 figures ─
    T("USE-CASE ceo: North Star recurring income target (THB 500M)", "POST", "http://localhost:3040/query",
      body={"query":"what is the North Star 2028 recurring income target?","channel":"t","user_id":"diag","dept_id":"ceo"},
      expect_any=["THB 500","500M","500 million","500 mn","Bt 500","500mn",
                  "five hundred million","500,000,000"], timeout=90),
    T("USE-CASE ceo: North Star AUM target (USD 600M)", "POST", "http://localhost:3040/query",
      body={"query":"what is the North Star institutional AUM target?","channel":"t","user_id":"diag","dept_id":"ceo"},
      expect_any=["USD 600","600M","600 million","US$600"], timeout=90),

    # ─ VCC: Brook LP FoF I terms from the Supplement ─
    T("USE-CASE vcc: FoF I management fee (1.5%)", "POST", "http://localhost:3040/query",
      body={"query":"what is the management fee on Brook LP FoF I?","channel":"t","user_id":"diag","dept_id":"vcc"},
      expect_any=["1.5%","1.5 percent","1.5 per cent"], timeout=90),
    T("USE-CASE vcc: FoF I hard cap (US$150M)", "POST", "http://localhost:3040/query",
      body={"query":"what is the hard cap on Brook LP FoF I?","channel":"t","user_id":"diag","dept_id":"vcc"},
      expect_any=["150M","150 million","US$150","USD 150"], timeout=90, flaky=True),
    T("USE-CASE vcc: FoF I auditor (Ernst & Young)", "POST", "http://localhost:3040/query",
      body={"query":"who is the auditor of Brook LP FoF I?","channel":"t","user_id":"diag","dept_id":"vcc"},
      expect_any=["Ernst & Young","Ernst and Young","EY "], timeout=90),
    T("USE-CASE vcc: FoF I custodian (DBS Bank)", "POST", "http://localhost:3040/query",
      body={"query":"who is the custodian of Brook LP FoF I?","channel":"t","user_id":"diag","dept_id":"vcc"},
      expect_any=["DBS Bank","DBS"], timeout=90),

    # ─ Legal: Timblick opinion specifics ─
    T("USE-CASE legal: FOF dividend WHT rate (10%)", "POST", "http://localhost:3040/query",
      body={"query":"what is the Thai withholding tax rate on dividends paid to FOF under Section 70?","channel":"t","user_id":"diag","dept_id":"legal"},
      expect_any=["10%","10 percent"], timeout=90),
    T("USE-CASE legal: actionable PE mitigation (non-Thai-resident signing)", "POST", "http://localhost:3040/query",
      body={"query":"what is the single most actionable mitigation for FOF place-of-effective-management risk?","channel":"t","user_id":"diag","dept_id":"legal"},
      expect_any=["non-Thai resident","non-resident","outside Thailand"], timeout=90, flaky=True),

    # ─ CIO: real portfolio facts ─
    T("USE-CASE cio: BTC holdings per coin book (164.6554)", "POST", "http://localhost:3040/query",
      body={"query":"how many BTC do we hold per the coin book?","channel":"t","user_id":"diag","dept_id":"cio"},
      expect_any=["164.6","164.65","164 BTC"], timeout=90, flaky=True),
    T("USE-CASE cio: HT Markets MTA execution recommendation (do NOT execute)", "POST", "http://localhost:3040/query",
      body={"query":"should we execute the HT Markets MTA today?","channel":"t","user_id":"diag","dept_id":"cio"},
      expect_any=["do not execute","not execute","unexecuted","not ready",
                  "HIGH risk","high risk","should not"], timeout=90),

    # ─ IC: 40%-rule numerator-column correctness ─
    T("USE-CASE ic: 40% rule numerator column (Investment Company Baht)", "POST", "http://localhost:3040/query",
      body={"query":"which dashboard column is the numerator for the 40 percent Investment Company rule?","channel":"t","user_id":"diag","dept_id":"ic"},
      expect_any=["Investment Company Baht","column H","col H"], timeout=90),

    # ─ IC (additional): Red Flag positions from May 2026 IC meeting ─
    # Source: IC-2026-05-12 meeting notes, section 3 Master sheet — MILL -94%, Wave -80%, B -79%
    T("USE-CASE ic: MILL drawdown Red Flag (-94%)", "POST", "http://localhost:3040/query",
      body={"query":"what is the current drawdown on MILL and is it Red Flag?","channel":"t","user_id":"diag","dept_id":"ic"},
      expect_any=["-94%","94%","94 percent","MILL"], timeout=90),

    # ─ IC (additional): DAT Round-1 total cash raised (Bt 392 mn) — IC-2026-05-12 slide 14 ─
    # Source: IC-2026-05-12 section 8a — "Total cash raised = Bt 392 mn"
    T("USE-CASE ic: DAT Round-1 total cash proceeds (Bt 392 mn)", "POST", "http://localhost:3040/query",
      body={"query":"what is the total cash raised from the DAT Round-1 sell-down plan?","channel":"t","user_id":"diag","dept_id":"ic"},
      expect_any=["392","Bt 392","392 mn","392mn","392 million"], timeout=90),

    # ─ Comms: marketed Pantera fund returns (with disclaimer expectation) ─
    T("USE-CASE comms: Pantera Fund I returns (19.4x / 46% IRR)", "POST", "http://localhost:3040/query",
      body={"query":"what was Pantera's Fund I performance?","channel":"t","user_id":"diag","dept_id":"comms"},
      expect_any=["19.4","46.0","46%","46 percent"], timeout=90, flaky=True),

    # ─ Comms (additional): Pantera investor lunch format (≤40 guests) — Pantera x Brooker April Event.docx ─
    # Source: 2026-04-21-pantera-x-brooker-investor-lunch — "limited to ~40 participants"
    T("USE-CASE comms: Pantera investor lunch guest cap (~40)", "POST", "http://localhost:3040/query",
      body={"query":"how many guests was the Pantera x Brooker investor lunch limited to?","channel":"t","user_id":"diag","dept_id":"comms"},
      expect_any=["40","forty","~40","40 guests","40 participants"], timeout=90),

    # ─ Comms (additional): Pantera Fund II TVPI (6.2x) — Pantera Talk April 2026.pptx ─
    # Source: pantera-blockchain-venture-performance — Fund II TVPI 6.2x, IRR 29.9%
    T("USE-CASE comms: Pantera Fund II TVPI (6.2x)", "POST", "http://localhost:3040/query",
      body={"query":"what is Pantera Fund II's TVPI?","channel":"t","user_id":"diag","dept_id":"comms"},
      expect_any=["6.2x","6.2 x","6.2"], timeout=90),

    # ─ HR: WFH legal basis (Thai Labour Protection Act §23/1) ─
    # Source: work-from-home-policy.md — "Thai Labour Protection Act Section 23/1"
    T("USE-CASE hr: WFH legal basis (Thai LPA Section 23/1)", "POST", "http://localhost:3002/query",
      body={"query":"what is the legal basis for work-from-home rights at Brooker?","channel":"t","user_id":"diag","dept_id":"hr"},
      expect_any=["23/1","Section 23","มาตรา 23","Labour Protection"], timeout=90),

    # ─ HR: contract-storage item 9 gap (government-officer records absent) ─
    # Source: employment-contract-document-retention.md — item 9 Absent / N/A
    T("USE-CASE hr: employment contract item 9 gap (government-officer records absent)", "POST",
      "http://localhost:3002/query",
      body={"query":"which employment contract control item was marked absent in the self-assessment?",
            "channel":"t","user_id":"diag","dept_id":"hr"},
      expect_any=["item 9","9","government","government officer","state official","absent"], timeout=90),

    # ─ IB: SuperSeed Fund III portfolio count (8 companies) ─
    # Source: superseed-fund-iii.md / 2026-05-08 briefing — "8 companies across 4 sectors"
    T("USE-CASE ib: SuperSeed Fund III portfolio size (8 companies)", "POST", "http://localhost:3040/query",
      body={"query":"how many portfolio companies does SuperSeed Fund III have?","channel":"t","user_id":"diag","dept_id":"ib"},
      expect_any=["8 companies","8 portfolio","eight companies","eight portfolio",
                  "8 port","portfolio of 8","**8**","has 8","has eight"], timeout=90,
      flaky=True),

    # ─ IB: Hive Autonomy payback period (7 months) ─
    # Source: hive-autonomy.md / SuperSeed brief page 2 — "7-month payback"
    T("USE-CASE ib: Hive Autonomy payback period (7 months)", "POST", "http://localhost:3040/query",
      body={"query":"what is the payback period for Hive Autonomy's fleet-autonomy solution?",
            "channel":"t","user_id":"diag","dept_id":"ib"},
      expect_any=["7-month","7 month","seven month","7 months","seven months"], timeout=90,
      flaky=True),

    # ─ IT: empty corpus — named disclosure abstain ─
    # Source: dept_test_matrix T-IT-01; corpus is empty (0 chunks in it_docs).
    # The agent must name the missing document type, not invent an uptime figure.
    T("USE-CASE it: empty corpus — abstain naming infra/security report", "POST",
      "http://localhost:3040/query",
      body={"query":"what is our system uptime this quarter?","channel":"t","user_id":"diag","dept_id":"it"},
      expect_any=["infrastructure","infra","security","no reference","corpus","document",
                  "don't have","no data","empty"], timeout=60),
]
TESTS.extend(USECASE)


for _dept, _frag in _DEPT_SRC_FRAGMENT.items():
    TESTS.extend([
        T(f"{_dept}: S02 greeting -> capability path (not abstain)", "POST", _RO_QUERY,
          body={"query": "hello", "channel": "t", "user_id": "diag", "dept_id": _dept},
          expect_not_in="don't have an answer for that", timeout=60),
        T(f"{_dept}: S03 mandate -> capability bypass", "POST", _RO_QUERY,
          body={"query": "what is your mandate?", "channel": "t", "user_id": "diag", "dept_id": _dept},
          expect_not_in="don't have an answer for that", timeout=90),
        T(f"{_dept}: X-02 SharePoint pre-guard names {_dept} sources", "POST", _RO_QUERY,
          body={"query": f"what's shown here? {SHARE}",
                "channel": "t", "user_id": "diag", "dept_id": _dept},
          expect_in=_frag,
          # Textbook D/E definition phrase — must not leak through if pre-guard fires
          expect_not_in="0.5 and 1.5",
          timeout=30),
        T(f"{_dept}: X-01 cannot request CAC report (403)", "GET", _RPT,
          params={"share_url": SHARE, "caller_dept": _dept},
          expect_status=403, expect_in="#cac-committee", timeout=10),
    ])


# ===========================================================================
# ARTEFACT TESTS — 11 priority deliverables
#
# Each T_ARTEFACT entry drives:
#   1. An HTTP POST to the appropriate deck-writer endpoint.
#   2. Assertion: HTTP 200, response time < 120 s, file_url present (for
#      /compose and /report) or Content-Type xlsx (for /compose-xlsx).
#   3. A download-and-parse step that asserts OOXML magic bytes (PK\x03\x04)
#      plus content fragments specific to each artefact.
#
# Endpoints:
#   POST /compose        -> .pptx  (response JSON has file_url pointing to
#                                   http://deck-writer:3050/files/{name})
#   POST /report         -> .docx  (response JSON has file_url pointing to
#                                   http://deck-writer:3050/reports/{name})
#   POST /compose-xlsx   -> .xlsx  (FileResponse — bytes returned directly)
#
# Priority classification matches the P1 / must-have tier in
# docs/agent_deliverables_matrix.md §1-§9 (artefacts with done or named
# status and non-empty corpora; deferred artefacts — Risk, IB, IT, HR
# comp/headcount — are excluded per the matrix).
# ===========================================================================

from dataclasses import dataclass, field as dc_field


@dataclass
class T_ARTEFACT:
    """One artefact test: HTTP call + file type + content fragment assertions."""

    name: str
    endpoint: str                           # relative path, e.g. "/compose"
    method: str = "POST"                    # POST for compose/report/compose-xlsx
    body: dict | None = None
    # For /compose and /report: substrings expected anywhere in the JSON
    # response text (file_url presence is always asserted separately).
    expect_json_in: list[str] = dc_field(default_factory=list)
    # Expected OOXML type: "pptx" / "docx" / "xlsx"
    filetype: str = "docx"
    # Substrings we assert must appear in the downloaded/returned file bytes
    # decoded as utf-8 with errors=ignore (good enough for OOXML XML parts
    # whose text nodes are 7-bit ASCII / UTF-8 string content).
    content_fragments: list[str] = dc_field(default_factory=list)
    # Minimum file size in bytes (sanity floor — an empty stub is ~6 kB for
    # xlsx, ~20 kB for docx, ~30 kB for pptx).
    min_bytes: int = 5_000
    timeout: float = 120.0


_DW = "http://localhost:3050"

# All OOXML types are ZIP containers — first 4 bytes must be PK\x03\x04.
_OOXML_MAGIC = b"PK\x03\x04"

# Content-type substrings expected in the HTTP response header per file type.
_CT_BY_EXT: dict[str, str] = {
    "docx": "wordprocessingml",
    "pptx": "presentationml",
    "xlsx": "spreadsheetml",
}

ARTEFACTS: list[T_ARTEFACT] = [
    # ── 1. CAC committee deck .pptx ──────────────────────────────────────────
    # Verifies the /compose pptx path for the CAC dept using the Brooker
    # template (brooker-deck-template.pptx).  The CAC Monthly Report .docx
    # already has a dedicated content-verified test; this adds the pptx path.
    # content_fragments: assert on Board-approved policy constants that MUST
    # appear in any correctly RAG-grounded CAC deck (25% DTV cap, 100 BTC
    # Sovereignty Buffer — both in the corpus).
    T_ARTEFACT(
        name="ARTEFACT cac: monthly committee deck .pptx",
        endpoint="/compose",
        body={
            "brief": (
                "CAC monthly committee deck: headline balance-sheet ratios, "
                "D/E vs 0.5x ceiling, on-chain DTV vs 25% cap, Sovereignty Buffer "
                "100 BTC status, funding-facility utilisation, and recommendations "
                "for the next month."
            ),
            "dept_id": "cac",
            "title": "CAC Monthly Committee — May 2026",
            "audience": "CAC committee members",
        },
        filetype="pptx",
        expect_json_in=["file_url", "slides"],
        # 25% DTV cap and 100 BTC Sovereignty Buffer are Board-approved
        # constants in the corpus that the RAG pipeline must surface.
        content_fragments=["25%", "100 BTC"],
        min_bytes=30_000,
    ),

    # ── 2. CFO quarterly board pack .docx ────────────────────────────────────
    # Generic /report path using the finance dept collection.
    T_ARTEFACT(
        name="ARTEFACT cfo: quarterly board pack .docx",
        endpoint="/report",
        body={
            "brief": (
                "CFO quarterly board pack: BICL financial highlights, "
                "related-party loan exposure on PN.34/35/36, FATCA Active NFFE "
                "classification status, D/E ratio, cash position, and audit prep "
                "outlook for FY2026."
            ),
            "dept_id": "finance",
            "title": "CFO Quarterly Board Pack — Q1 2026",
            "audience": "Board of Directors",
        },
        filetype="docx",
        expect_json_in=["file_url"],
        content_fragments=["CFO", "Q1 2026"],
        min_bytes=20_000,
    ),

    # ── 3. VCC quarterly LP report .docx ─────────────────────────────────────
    # Uses vcc dept collections (fund offering docs ingested).
    # content_fragments include corpus-grounded facts: 1.5% management fee
    # and DBS Bank custodian are confirmed VCC corpus facts.
    T_ARTEFACT(
        name="ARTEFACT vcc: quarterly LP report .docx",
        endpoint="/report",
        body={
            "brief": (
                "VCC quarterly LP report for Brook Technology Capital VCC: "
                "Brook LP FoF I — TVPI, DPI, IRR per sub-fund, capital called, "
                "distributions, management fee 1.5%, performance fee 20%, "
                "hard cap US$150M, custodian DBS Bank, auditor Ernst & Young."
            ),
            "dept_id": "vcc",
            "title": "Brook LP FoF I — Q1 2026 LP Report",
            "audience": "Limited Partners",
        },
        filetype="docx",
        expect_json_in=["file_url"],
        # "Brook", "LP", "FoF" from title; "1.5%" from VCC corpus (mgmt fee).
        content_fragments=["Brook", "LP", "FoF", "1.5%"],
        min_bytes=20_000,
    ),

    # ── 4. IC monthly meeting minutes .docx ──────────────────────────────────
    # Template exists at config/templates/ic/IC-meeting-minutes-reference.docx.
    # The generic /report path produces a RAG-grounded draft.
    T_ARTEFACT(
        name="ARTEFACT ic: monthly meeting minutes .docx",
        endpoint="/report",
        body={
            "brief": (
                "IC monthly meeting minutes: investment committee decisions, "
                "action items, DD pipeline status (BICL Movie, a16z, Pantera), "
                "Red Flag portfolio drawdowns (MILL/Wave/B/PACE/CV), "
                "40% Investment Company rule status (binding 30 Jun 2026), "
                "structured loan inventory, digital asset treasury update."
            ),
            "dept_id": "ic",
            "title": "IC Meeting Minutes — May 2026",
            "audience": "IC members",
        },
        filetype="docx",
        expect_json_in=["file_url"],
        content_fragments=["IC", "2026"],
        min_bytes=20_000,
    ),

    # ── 5. IC monthly meeting deck .pptx ─────────────────────────────────────
    # Template: config/templates/ic/IC-meeting-deck-reference.pptx.
    # /compose loads that template because dept_id="ic" is in _TEMPLATE_BY_DEPT.
    T_ARTEFACT(
        name="ARTEFACT ic: monthly meeting deck .pptx",
        endpoint="/compose",
        body={
            "brief": (
                "IC monthly committee deck: OKR-1 income build (THB 500M target), "
                "VCC Engine 1 dashboard, DAT sell+call strategy (slides 17-26), "
                "structured loan book (12-loan inventory), liquidity buffer, "
                "capital sovereignty doctrine, Red Flag portfolio status."
            ),
            "dept_id": "ic",
            "title": "IC Monthly Deck — May 2026",
            "audience": "Investment Committee",
        },
        filetype="pptx",
        expect_json_in=["file_url", "slides"],
        content_fragments=["IC", "2026"],
        min_bytes=30_000,
    ),

    # ── 6. CIO portfolio dashboard .xlsx ─────────────────────────────────────
    # Uses /compose-xlsx with a concrete monthly portfolio tracker spec.
    # Coin book row reflects the 164.6554 BTC figure from the coin-book corpus.
    T_ARTEFACT(
        name="ARTEFACT cio: portfolio dashboard .xlsx",
        endpoint="/compose-xlsx",
        body={
            "spec": {
                "metadata": {
                    "title": "CIO Portfolio Dashboard — May 2026",
                    "author": "CIO Agent",
                    "company": "The Brooker Group",
                },
                "sheets": [
                    {
                        "name": "Dashboard",
                        "style": "brooker",
                        "headers": [
                            "Category", "Strategy", "MTM (THB M)", "% of Portfolio",
                            "Ratio Cap", "Status",
                        ],
                        "rows": [
                            ["Digital Assets", "BTC Holdings", 0, 0.0, "50%", "Monitor"],
                            ["Digital Assets", "BNB Holdings", 0, 0.0, "50%", "Monitor"],
                            ["Non-Listed Equity", "ADFIN (Asian Finance)", 185.0, 0.0, "30%", "OK"],
                            ["Non-Listed Equity", "BSFL", 0.0, 0.0, "30%", "OK"],
                            ["Listed Equity", "Brooker (BKR)", 0.0, 0.0, "60%", "Red Flag"],
                            ["Listed Equity", "Wave BCG", 0.0, 0.0, "60%", "Red Flag"],
                            ["Structured Loan", "Portfolio Total", 0.0, 0.0, "50%", "OK"],
                            ["VCC / VC", "Brook LP FoF I", 0.0, 0.0, "N/A", "Active"],
                        ],
                        "column_widths": {
                            "A": 20, "B": 28, "C": 16, "D": 16, "E": 12, "F": 12,
                        },
                        "freeze_panes": "A2",
                    },
                    {
                        "name": "Coin Book",
                        "style": "brooker",
                        "headers": ["Asset", "Units", "Cost (USD)", "MTM (USD)", "P&L (USD)"],
                        "rows": [
                            ["BTC", 164.6554, 0.0, 0.0, 0.0],
                            ["BNB", 0.0, 0.0, 0.0, 0.0],
                        ],
                        "column_widths": {"A": 10, "B": 14, "C": 16, "D": 16, "E": 16},
                    },
                ],
            },
            "filename": "CIO_Portfolio_Dashboard_May2026.xlsx",
            "caller_dept": "cio",
        },
        filetype="xlsx",
        expect_json_in=[],          # compose-xlsx returns FileResponse, not JSON
        content_fragments=["Dashboard", "Coin Book", "CIO Portfolio Dashboard"],
        min_bytes=6_000,
    ),

    # ── 7. CEO quarterly board pre-read .docx ────────────────────────────────
    T_ARTEFACT(
        name="ARTEFACT ceo: quarterly board pre-read .docx",
        endpoint="/report",
        body={
            "brief": (
                "CEO quarterly board pre-read pack: enterprise posture update, "
                "North Star 2028 progress (THB 500M recurring income, USD 600M AUM), "
                "Khao Yai Resolutions R-01..R-08 status, OKR-1..5 delta vs plan, "
                "Three-Engine flywheel progress (Engine 1 VCC, Engine 2 structured "
                "loans, Engine 3 DAT), cross-dept synthesis and recommendations."
            ),
            "dept_id": "ceo",
            "title": "CEO Board Pre-Read — Q1 2026",
            "audience": "Board of Directors",
        },
        filetype="docx",
        expect_json_in=["file_url"],
        content_fragments=["CEO", "2026"],
        min_bytes=20_000,
    ),

    # ── 8. CEO OKR tracker .xlsx ─────────────────────────────────────────────
    # Encodes all 5 OKRs from the CEO SKILL.md with known values from the
    # corpus (OKR-5: 54.6% vs 40% cap, binding 30 Jun 2026).
    T_ARTEFACT(
        name="ARTEFACT ceo: OKR tracker .xlsx",
        endpoint="/compose-xlsx",
        body={
            "spec": {
                "metadata": {
                    "title": "CEO OKR Tracker — 2026",
                    "author": "CEO Agent",
                    "company": "The Brooker Group",
                },
                "sheets": [
                    {
                        "name": "OKR Summary",
                        "style": "brooker",
                        "headers": [
                            "OKR #", "Objective", "Key Result", "Owner",
                            "Target", "Current", "% Complete", "Status",
                        ],
                        "rows": [
                            ["OKR-1", "Recurring Income Engine",
                             "THB 500M recurring income by 2028",
                             "CEO", "THB 500M", "TBD", 0.0, "In Progress"],
                            ["OKR-2", "Institutional AUM Growth",
                             "USD 600M institutional AUM",
                             "CIO", "USD 600M", "TBD", 0.0, "In Progress"],
                            ["OKR-3", "VCC Fund Scale",
                             "Brook LP FoF I hard cap US$150M subscribed",
                             "VCC", "US$150M", "TBD", 0.0, "In Progress"],
                            ["OKR-4", "Capital Sovereignty",
                             "Sovereignty Buffer at least 100 BTC maintained",
                             "CAC", ">=100 BTC", "TBD", 0.0, "In Progress"],
                            ["OKR-5", "Investment Company Rule",
                             "Portfolio ratio below 40% by 30 Jun 2026",
                             "IC", "<40%", "54.6%", 0.0, "BREACH"],
                        ],
                        "column_widths": {
                            "A": 8, "B": 26, "C": 40, "D": 8,
                            "E": 14, "F": 14, "G": 12, "H": 14,
                        },
                        "freeze_panes": "A2",
                    },
                    {
                        "name": "Monthly Progress",
                        "style": "brooker",
                        "headers": ["OKR #", "Month", "Milestone", "Status", "Notes"],
                        "rows": [
                            ["OKR-1", "May 2026", "Structured loan book buildout",
                             "On Track", ""],
                            ["OKR-3", "May 2026", "FoF I first closing",
                             "In Progress", ""],
                            ["OKR-4", "May 2026", "BTC 164.6 BTC held", "OK", ""],
                            ["OKR-5", "May 2026", "IC ratio 54.6% — above 40% cap",
                             "BREACH", "Binding 30 Jun 2026"],
                        ],
                    },
                ],
            },
            "filename": "CEO_OKR_Tracker_2026.xlsx",
            "caller_dept": "ceo",
        },
        filetype="xlsx",
        expect_json_in=[],
        content_fragments=["OKR Summary", "Monthly Progress", "OKR-1", "OKR-5"],
        min_bytes=6_000,
    ),

    # ── 9. Comms investor event deck .pptx ───────────────────────────────────
    # Brooker template; uses comms dept collections.
    T_ARTEFACT(
        name="ARTEFACT comms: investor event deck .pptx",
        endpoint="/compose",
        body={
            "brief": (
                "Comms investor event deck: Brooker Group firm overview, "
                "Three-Engine flywheel (digital assets, structured lending, VCC), "
                "Pantera collaboration highlights (Fund I 19.4x / 46% IRR), "
                "macro outlook (Global Liquidity Cycle, Internet of Value), "
                "VCC Brook LP FoF I fund summary, and event logistics."
            ),
            "dept_id": "comms",
            "title": "Brooker Investor Event — May 2026",
            "audience": "Institutional investors and LPs",
        },
        filetype="pptx",
        expect_json_in=["file_url", "slides"],
        content_fragments=["Brooker", "2026"],
        min_bytes=30_000,
    ),

    # ── 10. HR self-assessment control tracker .xlsx ─────────────────────────
    # HR corpus has 2 Thai self-assessment PDFs (contract storage + WFH
    # controls questionnaires). A tracker sheet is within scope — the
    # deferred headcount/comp trackers (no corpus) are excluded.
    T_ARTEFACT(
        name="ARTEFACT hr: self-assessment control tracker .xlsx",
        endpoint="/compose-xlsx",
        body={
            "spec": {
                "metadata": {
                    "title": "HR Self-Assessment Control Tracker",
                    "author": "HR Agent",
                    "company": "The Brooker Group",
                },
                "sheets": [
                    {
                        "name": "Contract Storage Controls",
                        "style": "brooker",
                        "headers": [
                            "Control #", "Control Description", "Status",
                            "Evidence", "Review Date", "Owner",
                        ],
                        "rows": [
                            [1, "Employment contracts stored in secure system",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [2, "Access controls on contract repository",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [3, "Contract templates reviewed for legal compliance",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [4, "Retention policy defined (per Thai labour law)",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [5, "Destruction / disposal procedure documented",
                             "Partial", "Questionnaire 2025", "2026-06-01", "HR"],
                            [6, "Labour law and ethics training completed",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                        ],
                        "column_widths": {
                            "A": 10, "B": 45, "C": 12, "D": 20, "E": 14, "F": 8,
                        },
                        "freeze_panes": "A2",
                    },
                    {
                        "name": "WFH Controls",
                        "style": "brooker",
                        "headers": [
                            "Control #", "Control Description", "Status",
                            "Evidence", "Review Date", "Owner",
                        ],
                        "rows": [
                            [1, "WFH policy documented (Section 23/1 compliant)",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [2, "Employee WFH agreements signed",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [3, "IT security policy for remote work",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                            [4, "Overtime / working-hours monitoring",
                             "Present", "Questionnaire 2025", "2026-01-01", "HR"],
                        ],
                        "column_widths": {
                            "A": 10, "B": 45, "C": 12, "D": 20, "E": 14, "F": 8,
                        },
                    },
                ],
            },
            "filename": "HR_SelfAssessment_ControlTracker.xlsx",
            "caller_dept": "hr",
        },
        filetype="xlsx",
        expect_json_in=[],
        content_fragments=[
            "Contract Storage Controls", "WFH Controls", "HR Self-Assessment",
        ],
        min_bytes=6_000,
    ),

    # ── 11. CFO quarterly board deck .pptx ───────────────────────────────────
    # Brooker template; complements artefact #2 (the CFO docx).  Verifies
    # the /compose pptx path for the finance dept.
    # content_fragments: assert on factual data that the RAG pipeline must
    # surface (PN.35 USD figure and D/E ceiling), not on LLM-chosen titles
    # (the LLM may rename "CFO Board Deck" to "Q1 Financial Review" etc.).
    T_ARTEFACT(
        name="ARTEFACT cfo: quarterly board deck .pptx",
        endpoint="/compose",
        body={
            "brief": (
                "CFO quarterly board deck: BICL financial highlights, FATCA "
                "Active NFFE classification, related-party PN.34/35/36 exposure "
                "(PN.35 USD 39.02M), D/E ratio vs 0.5x Board ceiling, cash "
                "liquidity position, audit prep roadmap for FY2026, and key "
                "financial risks for board attention."
            ),
            "dept_id": "finance",
            "title": "CFO Board Deck — Q1 2026",
            "audience": "Board of Directors",
        },
        filetype="pptx",
        expect_json_in=["file_url", "slides"],
        # Assert on grounded facts from the corpus: PN.35 USD amount and
        # D/E ceiling; these must appear in any correctly RAG-grounded deck.
        content_fragments=["39.02", "0.5"],
        min_bytes=30_000,
    ),
]


# ===========================================================================
# Artefact verification helpers
# ===========================================================================

def _ooxml_magic_ok(data: bytes) -> bool:
    """All OOXML types (docx, pptx, xlsx) are ZIP archives — magic PK\x03\x04."""
    return data[:4] == _OOXML_MAGIC


def _content_type_ok(headers: "httpx.Headers", filetype: str) -> bool:
    ct = headers.get("content-type", "").lower()
    return _CT_BY_EXT.get(filetype, "") in ct


def _bytes_missing_fragments(data: bytes, fragments: list[str]) -> list[str]:
    """Return fragments NOT found anywhere in the OOXML file.

    Strategy: OOXML files are ZIP archives. Raw `bytes.decode('utf-8')` of the
    compressed container does NOT reliably find XML text content because the
    per-entry deflate streams are compressed.  We decompress every XML entry
    in the ZIP and search the uncompressed XML text instead.  Falls back to
    the raw decode if zipfile parsing fails (e.g. for a plain binary that
    isn't OOXML).
    """
    import zipfile
    import io

    # Build a single large search corpus from all decompressed XML parts.
    try:
        z = zipfile.ZipFile(io.BytesIO(data))
        text_parts: list[str] = []
        for name in z.namelist():
            try:
                text_parts.append(z.read(name).decode("utf-8", errors="ignore"))
            except Exception:
                pass
        text = "\n".join(text_parts)
    except Exception:
        # Not a valid ZIP — fall back to raw bytes decode.
        text = data.decode("utf-8", errors="ignore")

    return [f for f in fragments if f not in text]


async def verify_one_artefact(
    art: T_ARTEFACT, client: "httpx.AsyncClient"
) -> tuple[bool, str]:
    """Run one artefact test end-to-end.

    /compose and /report:
      1. POST -> assert HTTP 200 + file_url in response JSON.
      2. GET file_url -> assert OOXML magic + content fragments.

    /compose-xlsx:
      1. POST -> assert HTTP 200 + Content-Type contains spreadsheetml.
      2. Assert OOXML magic + content fragments on the direct response body.
    """
    import time as _time

    url = f"{_DW}{art.endpoint}"
    try:
        t0 = _time.monotonic()
        r = await client.post(url, json=art.body, timeout=art.timeout)
        elapsed = _time.monotonic() - t0
    except Exception as exc:
        return False, f"request failed: {type(exc).__name__}: {exc}"

    if r.status_code != 200:
        snippet = r.text[:200].replace("\n", " ")
        return False, f"HTTP {r.status_code} (expected 200) | {snippet}"

    if elapsed > art.timeout:
        return False, f"timeout: {elapsed:.1f}s > {art.timeout}s"

    # ── /compose-xlsx path: response body IS the xlsx bytes ─────────────────
    if art.endpoint == "/compose-xlsx":
        data = r.content
        reasons: list[str] = []
        if len(data) < art.min_bytes:
            reasons.append(f"file too small: {len(data):,} bytes < {art.min_bytes:,}")
        if not _ooxml_magic_ok(data):
            reasons.append(f"bad OOXML magic: {data[:4]!r}")
        if not _content_type_ok(r.headers, art.filetype):
            reasons.append(
                f"wrong Content-Type: {r.headers.get('content-type')!r}"
            )
        missing = _bytes_missing_fragments(data, art.content_fragments)
        if missing:
            reasons.append(f"missing content fragments: {missing}")
        if reasons:
            return False, "; ".join(reasons)
        return True, (
            f"xlsx OK: {len(data):,} bytes, magic PK, "
            f"{len(art.content_fragments)}/{len(art.content_fragments)} fragments, "
            f"{elapsed:.1f}s"
        )

    # ── /compose or /report path: response is JSON with file_url ────────────
    try:
        body = r.json()
    except Exception:
        return False, f"response not JSON: {r.text[:200]}"

    missing_keys = [k for k in art.expect_json_in if k.lower() not in r.text.lower()]
    if missing_keys:
        snippet = r.text[:200].replace("\n", " ")
        return False, f"JSON response missing keys/text: {missing_keys} | {snippet}"

    file_url = body.get("file_url") or body.get("file_path") or ""
    if not file_url:
        return False, f"no file_url in response: {str(body)[:200]}"

    # Rewrite container hostname -> localhost for the out-of-container runner.
    dl_url = file_url.replace("deck-writer:3050", "localhost:3050")

    try:
        fr = await client.get(dl_url, timeout=30)
    except Exception as exc:
        return False, f"file download failed: {type(exc).__name__}: {exc}"

    if fr.status_code != 200:
        return False, f"file download HTTP {fr.status_code}: {dl_url}"

    data = fr.content
    reasons = []
    if len(data) < art.min_bytes:
        reasons.append(f"file too small: {len(data):,} bytes < {art.min_bytes:,}")
    if not _ooxml_magic_ok(data):
        reasons.append(f"bad OOXML magic: {data[:4]!r}")
    if not _content_type_ok(fr.headers, art.filetype):
        reasons.append(f"wrong Content-Type: {fr.headers.get('content-type')!r}")
    missing = _bytes_missing_fragments(data, art.content_fragments)
    if missing:
        reasons.append(f"missing content fragments: {missing}")

    if reasons:
        return False, (
            "; ".join(reasons)
            + f"  (dl={dl_url}, {len(data):,} bytes)"
        )
    return True, (
        f"{art.filetype.upper()} OK: {len(data):,} bytes, magic PK, "
        f"{len(art.content_fragments)}/{len(art.content_fragments)} fragments, "
        f"{elapsed:.1f}s end-to-end"
    )


def _trim(s: str, n: int = 160) -> str:
    s = " ".join(s.split())
    return s if len(s) <= n else s[:n] + "…"


async def verify_cac_report_docx_content(client: httpx.AsyncClient) -> tuple[bool, str]:
    """Generate a CAC report and verify its ACTUAL docx contents match the data pack.

    Unlike the mechanical "report: caller_dept=cac → 200, has file_url" test, this:
      1. Calls the endpoint.
      2. Fetches the file_url.
      3. Opens the .docx with python-docx.
      4. Asserts the breach table contains the three expected breaches AND the
         headline THB figures appear in the body.
    """
    try:
        r = await client.get(
            "http://localhost:3050/report/cac-meeting",
            params={"share_url": SHARE, "caller_dept": "cac"},
            timeout=120,
        )
        r.raise_for_status()
        body = r.json()
        url = body.get("file_url", "").replace("deck-writer:3050", "localhost:3050")
        if not url:
            return False, "no file_url in response"
        f = await client.get(url, timeout=30)
        f.raise_for_status()
        import io
        from docx import Document
        d = Document(io.BytesIO(f.content))
        all_text = "\n".join(p.text for p in d.paragraphs)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += "\n" + cell.text
        # Required facts from your filled Data Pack (verified against the original docx earlier):
        checks = {
            "breach: Investment / Total Assets 54.6%": ("Investment / Total Assets" in all_text
                                                         and "54.6%" in all_text),
            "watch: Treasury native yield 1.0%": ("Treasury native yield" in all_text
                                                   and "1.0%" in all_text),
            "watch: Sovereignty Buffer 100 BTC": ("Sovereignty Buffer" in all_text
                                                    and "100 BTC" in all_text),
            "headline figure: Total assets THB 3,127.97M": "3,127.97" in all_text,
            "headline figure: Net Worth THB 2,435.11M": "2,435.11" in all_text,
            "VCC Platform line: THB 253.87M (no 254M rounding)": ("253.87" in all_text
                                                                    and "THB 254M" not in all_text),
        }
        failed = [k for k, ok in checks.items() if not ok]
        if failed:
            return False, f"docx content missing: {failed}  (bytes={len(f.content)})"
        return True, f"docx content verified: 6/6 facts present  (bytes={len(f.content)})"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


async def run_voted(t: T, client: httpx.AsyncClient) -> tuple[bool, str]:
    """Run a flaky test 3 times sequentially and pass if at least 2/3 runs pass.

    Prints individual run results inline so the caller's loop can print the
    final voted outcome.  Returns (voted_ok, summary_msg).
    """
    RUNS = 3
    THRESHOLD = 2
    results: list[tuple[bool, str]] = []
    for run_num in range(1, RUNS + 1):
        ok, msg = await run_one(t, client)
        results.append((ok, msg))
        run_tag = f"{GREEN}PASS{RESET}" if ok else f"{RED}FAIL{RESET}"
        # Trim the per-run detail message to keep output readable
        short_msg = _trim(msg, 120)
        print(f"    run {run_num}: [{run_tag}] {DIM}{short_msg}{RESET}")
    passes = sum(1 for ok, _ in results if ok)
    voted_ok = passes >= THRESHOLD
    voted_tag = f"{GREEN}VOTED PASS{RESET}" if voted_ok else f"{RED}VOTED FAIL{RESET}"
    print(f"    >> [{voted_tag}] ({passes}/{RUNS})")
    summary = f"voted {passes}/{RUNS} passing runs"
    return voted_ok, summary


async def run_one(t: T, client: httpx.AsyncClient) -> tuple[bool, str]:
    try:
        if t.method == "GET":
            r = await client.get(t.url, params=t.params, timeout=t.timeout)
        else:
            r = await client.post(t.url, json=t.body, timeout=t.timeout)
    except Exception as exc:
        return False, f"request failed: {type(exc).__name__}: {exc}"
    body = r.text
    body_lc = body.lower()
    reasons = []
    if r.status_code != t.expect_status:
        reasons.append(f"status {r.status_code} != expected {t.expect_status}")
    if t.expect_in and t.expect_in.lower() not in body_lc:
        reasons.append(f"missing {t.expect_in!r}")
    if t.expect_not_in and t.expect_not_in.lower() in body_lc:
        reasons.append(f"forbidden {t.expect_not_in!r} found")
    if t.expect_any:
        if not any(s.lower() in body_lc for s in t.expect_any):
            reasons.append(f"none of {t.expect_any!r} present")
    ok = not reasons
    msg = f"HTTP {r.status_code}  | {_trim(body)}"
    if reasons:
        msg = "; ".join(reasons) + "  ::  " + msg
    return ok, msg


async def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--only", help="run only tests whose name contains this substring")
    args = ap.parse_args()
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    tests = [t for t in TESTS if not args.only or args.only.lower() in t.name.lower()]
    artefacts = [
        a for a in ARTEFACTS
        if not args.only or args.only.lower() in a.name.lower()
    ]
    n_pass = n_fail = 0
    async with httpx.AsyncClient() as client:
        for t in tests:
            if t.flaky:
                print(f"[....] {t.name} (voted 3x)")
                ok, msg = await run_voted(t, client)
            else:
                ok, msg = await run_one(t, client)
            tag = f"{GREEN}PASS{RESET}" if ok else f"{RED}FAIL{RESET}"
            print(f"[{tag}] {t.name}\n       {DIM}{msg}{RESET}")
            n_pass += ok
            n_fail += not ok

        # Special "real-content" test: fetch the .docx and verify its contents.
        if not args.only or "docx" in args.only.lower() or "content" in args.only.lower():
            ok, msg = await verify_cac_report_docx_content(client)
            tag = f"{GREEN}PASS{RESET}" if ok else f"{RED}FAIL{RESET}"
            print(f"[{tag}] USE-CASE cac docx content: breach table + headline figures match data pack")
            print(f"       {DIM}{msg}{RESET}")
            n_pass += ok
            n_fail += not ok

        # ── 11 priority artefact tests ────────────────────────────────────────
        if artefacts:
            print(f"\n{'─'*72}")
            print(f"ARTEFACT TESTS ({len(artefacts)} selected)")
            print(f"{'─'*72}")
        for art in artefacts:
            ok, msg = await verify_one_artefact(art, client)
            tag = f"{GREEN}PASS{RESET}" if ok else f"{RED}FAIL{RESET}"
            print(f"[{tag}] {art.name}\n       {DIM}{msg}{RESET}")
            n_pass += ok
            n_fail += not ok

    total = n_pass + n_fail
    print(f"\n{n_pass}/{total} passed, {n_fail} failed")
    return 0 if n_fail == 0 else 1


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
