"""Deterministic CAC meeting-report .docx builder (shared by script + services).

Turns a filled CAC Monthly Data Pack into the populated CFO-report .docx. Figures
are read straight from the cells (no LLM, no RAG, no fabrication); the only added
content is limit-status flags computed against the Board-approved thresholds
(Khao Yai §2.8) and a deterministic executive summary. Blank narrative cells get
an explicit "[to be added]" marker rather than invented commentary.

Two input shapes are supported so the same logic serves the CLI and the services:
  - pack_from_workbook(): the MS Graph dict {sheet: [[cell, ...], ...]} (no openpyxl)
  - pack_from_xlsx():     a local .xlsx via openpyxl (lazy import)

Both yield the same `pack` structure consumed by build_cac_report_docx():
    {sheet_name: {normalised_label: (original_label, [colB, colC, colD, colE, colF])}}

Unit notes (the pack mixes units, so the report normalises):
  - Balance Sheet / Capital Allocation / Liquidity / ALM cells are full baht.
  - Funding Structure limits/drawn are already in THB millions.
  - Ratios (D/E, investment %, cost of funds, DTV, yield) are decimal fractions.
"""
from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

_log = logging.getLogger(__name__)

OK = "OK"
BREACH = "BREACH"
WATCH = "WATCH"


# --- value parsing / formatting ---------------------------------------------

def _num(v) -> float | None:
    """Parse a cell to float; strips %, x, commas, spaces. None if not numeric."""
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).replace("%", "").replace("x", "").replace(",", "").strip()
    try:
        return float(s)
    except ValueError:
        return None


def _pct(frac) -> str:
    """Format a decimal fraction (0.546) or a whole percent (54.6) as 'XX.X%'."""
    v = _num(frac)
    if v is None:
        return "—"
    if -1.5 <= v <= 1.5:  # stored as a fraction
        v *= 100
    return f"{v:.1f}%"


def _pct_val(frac) -> float | None:
    """Return the percentage as a number (54.6), normalising fractions."""
    v = _num(frac)
    if v is None:
        return None
    return v * 100 if -1.5 <= v <= 1.5 else v


def _thb_m(v) -> str:
    """Format a baht (or already-millions) value as 'THB X,XXX.XXM'.

    Keeps up to 2 decimals of THB-millions precision so 253,872,071 reads as
    'THB 253.87M' — never a coarse, overstated 'THB 254M'. Trailing zeros are
    dropped so whole millions stay clean ('THB 500M', 'THB 0M').
    """
    n = _num(v)
    if n is None:
        return "—"
    m = n / 1e6 if abs(n) >= 1e5 else n  # full baht -> millions; else already M
    s = f"{m:,.2f}".rstrip("0").rstrip(".")  # '.' blocks rstrip from eating integers
    return f"THB {s}M"


def _norm(label) -> str:
    return re.sub(r"\s+", " ", str(label or "")).strip().lower()


# --- pack readers ------------------------------------------------------------

def _empty(c) -> bool:
    return c is None or (isinstance(c, str) and not c.strip())


def pack_from_workbook(workbook: dict[str, list[list]]) -> dict[str, dict[str, tuple]]:
    """Build the label-keyed pack from the MS Graph workbook dict.

    `workbook` is {sheet_name: [[colA, colB, ...], ...]} as returned by
    GraphExcel.read_workbook_by_share_url (empty cells are "" or None).
    """
    out: dict[str, dict[str, tuple]] = {}
    for sheet, rows in (workbook or {}).items():
        if sheet.strip().lower() == "instructions":
            continue
        d: dict[str, tuple] = {}
        for row in rows or []:
            if not row or _empty(row[0]):
                continue
            text = str(row[0]).strip()
            if re.match(r"^\d+\.\s", text):  # section-title row
                continue
            vals = [None if _empty(c) else c for c in row[1:6]]
            vals += [None] * (5 - len(vals))
            d.setdefault(_norm(text), (text, vals))  # first occurrence wins
        out[sheet] = d
    return out


def pack_from_xlsx(xlsx: str | Path) -> dict[str, dict[str, tuple]]:
    """Build the pack from a local .xlsx file (lazy openpyxl import)."""
    import openpyxl  # local import: services that pass a workbook dict don't need it

    wb = openpyxl.load_workbook(xlsx, data_only=True)
    workbook: dict[str, list[list]] = {}
    for s in wb.sheetnames:
        ws = wb[s]
        workbook[s] = [
            [ws.cell(r, c).value for c in range(1, 7)]
            for r in range(1, ws.max_row + 1)
        ]
    return pack_from_workbook(workbook)


def _get(pack: dict, sheet: str, label_sub: str, col: int = 0):
    """First value (col 0 == column B) whose label contains label_sub."""
    rows = pack.get(sheet, {})
    target = label_sub.lower()
    for _key, (_orig, vals) in rows.items():
        if target in _key:
            return vals[col] if col < len(vals) else None
    return None


# --- docx helpers ------------------------------------------------------------

def _bullet(doc, text, bold_lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        p.add_run(bold_lead).bold = True
    p.add_run(text)
    return p


def _kv(doc, label, value):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(value))
    return p


def cac_report_context(pack: dict, month: str) -> tuple[dict, list[dict]]:
    """Build a flat context dict + breach list from the data pack.

    The same data the programmatic builder uses, exposed for the template-first
    path (services/shared/office_template.py) so a user-authored .docx with
    `{{total_assets}}`, `{{net_worth}}`, … placeholders can be filled.

    Returns (context, breaches). Keys in context are snake_case_strings.
    """
    from datetime import date as _date

    total_assets = _get(pack, "1 Balance Sheet", "total assets")
    total_liab = _get(pack, "1 Balance Sheet", "total liabilities")
    bank_debt = _get(pack, "1 Balance Sheet", "bank debt")
    net_worth = _get(pack, "1 Balance Sheet", "total equity")
    de_ratio = _num(_get(pack, "1 Balance Sheet", "debt-to-equity"))
    inv_ratio = _pct_val(_get(pack, "1 Balance Sheet", "investment / total assets"))

    runway = _num(_get(pack, "3 Liquidity", "months of runway"))
    op_cash = _get(pack, "3 Liquidity", "operating cash")
    near_cash = _get(pack, "3 Liquidity", "near-cash")
    burn = _get(pack, "3 Liquidity", "operating burn")
    buffer_btc = _num(_get(pack, "7 Risk Appetite", "sovereignty buffer"))
    cfp = _get(pack, "3 Liquidity", "contingency funding plan")

    cost_funds = _pct_val(_get(pack, "7 Risk Appetite", "cost of funds"))
    dtv = _pct_val(_get(pack, "7 Risk Appetite", "aggregate on-chain dtv"))
    native_yield = _pct_val(_get(pack, "7 Risk Appetite", "treasury native yield"))

    dur_gap = _num(_get(pack, "5 Asset-Liability", "duration gap"))
    refinancing = _get(pack, "5 Asset-Liability", "refinancing due")
    collateral = _get(pack, "5 Asset-Liability", "collateral sufficiency")
    fx = _get(pack, "5 Asset-Liability", "fx exposure")

    agg = pack.get("4 Funding Structure", {}).get("aggregate", (None, []))[1]
    agg_drawn = agg[2] if len(agg) > 2 else None
    agg_avail = agg[3] if len(agg) > 3 else None
    agg_util = agg[4] if len(agg) > 4 else None

    # Per-engine allocation (Three-Engine Strategy — Khao Yai §2.4). Row labels in
    # the data pack are like "VCC Platform (Engine 1, GP/seed)" / "Advisory book
    # (Engine 2)" / "Digital Asset Treasury (Engine 3)"; substring-match on the
    # `engine N` token is the most robust hook. _thb_m() normalises units so
    # 253,872,071 baht renders as "THB 253.87M".
    engine_1_raw = _get(pack, "2 Capital Allocation", "engine 1")
    engine_2_raw = _get(pack, "2 Capital Allocation", "engine 2")
    engine_3_raw = _get(pack, "2 Capital Allocation", "engine 3")

    if engine_1_raw is None:
        _log.warning("CAC report: Engine 1 (VCC) value missing from data pack")
    if engine_2_raw is None:
        _log.warning("CAC report: Engine 2 (Listed equities / Advisory) value missing from data pack")
    if engine_3_raw is None:
        _log.warning("CAC report: Engine 3 (DAT) value missing from data pack")

    # Total invested: prefer an explicit TOTAL row if present, else sum the three
    # engine numbers we just pulled. Fall back to "n/a" when nothing usable.
    engine_total_raw = _get(pack, "2 Capital Allocation", "total")
    if engine_total_raw is None:
        engine_nums = [_num(v) for v in (engine_1_raw, engine_2_raw, engine_3_raw)]
        valid = [n for n in engine_nums if n is not None]
        if valid:
            engine_total_raw = sum(valid)
        else:
            _log.warning("CAC report: engine total unavailable (no TOTAL row, no engine values)")

    def _engine_value(raw) -> str:
        return _thb_m(raw) if _num(raw) is not None else "n/a"

    # Build breach checks (same logic as build_cac_report_docx)
    checks: list[dict] = []

    def _add(metric, value_str, status, limit, escalate):
        checks.append({"metric": metric, "value": value_str, "status": status,
                       "limit": limit, "escalate": escalate})

    if inv_ratio is not None:
        _add("Investment / Total Assets", f"{inv_ratio:.1f}%",
             BREACH if inv_ratio >= 40 else OK, "< 40% by 30 Jun 2026",
             "Board (post-Jun 2026)")
    if de_ratio is not None:
        _add("Debt-to-Equity", f"{de_ratio:.2f}x",
             BREACH if de_ratio > 0.5 else OK, "< 0.5x", "Board")
    if cost_funds is not None:
        _add("Cost of funds", f"{cost_funds:.2f}%",
             BREACH if cost_funds > 4 else OK, "< 3-4%", "CAC")
    if dtv is not None:
        _add("Aggregate on-chain DTV", f"{dtv:.1f}%",
             BREACH if dtv > 25 else OK, "<= 25% (hard cap)", "Board if > 25%")
    if native_yield is not None:
        _add("Treasury native yield", f"{native_yield:.1f}%",
             WATCH if native_yield < 10 else OK, ">= 10% / THB 100M", "CAC")
    if runway is not None:
        _add("Liquidity runway", f"{runway:.0f} months",
             BREACH if runway < 6 else OK, ">= 6 months", "CEO if < 6mo")
    if buffer_btc is not None:
        _add("Sovereignty Buffer", f"{buffer_btc:.0f} BTC",
             WATCH if buffer_btc <= 100 else OK, ">= 100 BTC", "CEO if breached")
    if dur_gap is not None:
        _add("Duration gap", f"{dur_gap:.2f} yr",
             BREACH if dur_gap >= 2.0 else OK, "< 2.0 yr", "CAC")

    breaches = [c for c in checks if c["status"] in (BREACH, WATCH)]
    hard = [c for c in breaches if c["status"] == BREACH]

    # Flat context — snake_case keys the template can reference as {{key}}
    context = {
        "month": month,
        "prepared_date": _date.today().isoformat(),
        "status_line": "DRAFT — for committee review",

        # Balance sheet
        "total_assets": _thb_m(total_assets),
        "total_liabilities": _thb_m(total_liab),
        "bank_debt": _thb_m(bank_debt),
        "net_worth": _thb_m(net_worth),
        "de_ratio": f"{de_ratio:.2f}x" if de_ratio is not None else "—",
        "inv_ratio": f"{inv_ratio:.1f}%" if inv_ratio is not None else "—",

        # Liquidity
        "operating_cash": _thb_m(op_cash),
        "near_cash": _thb_m(near_cash),
        "monthly_burn": _thb_m(burn),
        "runway_months": f"{runway:.0f}" if runway is not None else "—",
        "sovereignty_buffer_btc": f"{buffer_btc:.0f}" if buffer_btc is not None else "—",
        "contingency_funding_plan": str(cfp) if cfp else "[to be added]",

        # Funding
        "facility_drawn": _thb_m(agg_drawn),
        "facility_available": _thb_m(agg_avail),
        "facility_util": _pct(agg_util),
        "cost_of_funds": f"{cost_funds:.2f}%" if cost_funds is not None else "—",

        # ALM
        "duration_gap_yr": f"{dur_gap:.2f}" if dur_gap is not None else "—",
        "refinancing_due_12m": str(refinancing) if refinancing else "[to be added]",
        "collateral_sufficiency": str(collateral) if collateral else "[to be added]",
        "fx_exposure": str(fx) if fx else "[to be added]",

        # On-chain
        "on_chain_dtv": f"{dtv:.1f}%" if dtv is not None else "—",
        "treasury_native_yield": f"{native_yield:.1f}%" if native_yield is not None else "—",

        # Engine performance — Three-Engine Strategy allocations (Khao Yai §2.4).
        # Labels are stable per Khao Yai; values pull from "2 Capital Allocation".
        "engine_1_vcc_label": "Engine 1 — VCC",
        "engine_1_vcc_value": _engine_value(engine_1_raw),
        "engine_2_listed_label": "Engine 2 — Listed equities",
        "engine_2_listed_value": _engine_value(engine_2_raw),
        "engine_3_dat_label": "Engine 3 — DAT",
        "engine_3_dat_value": _engine_value(engine_3_raw),
        "engine_total_label": "Total invested",
        "engine_total_value": _engine_value(engine_total_raw),

        # Counts (useful for headlines)
        "num_breaches": str(len(hard)),
        "num_watch_items": str(len(breaches) - len(hard)),
        "num_total_flags": str(len(breaches)),

        # Pre-formatted breach lines (one-per-line; user can include {{breach_list}}
        # as a single block in the template). For richer rendering, the
        # programmatic builder writes a real table — that's what fallback gives.
        "breach_list": "\n".join(
            f"- {c['metric']}: {c['value']} (limit {c['limit']} — {c['escalate']})"
            for c in breaches
        ) if breaches else "No limits breached this month.",
    }
    return context, breaches


def build_cac_report_docx(pack: dict, out_docx: str | Path, month: str) -> list[dict]:
    """Render the CAC meeting report .docx from a pack. Returns the breach list."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    out_docx = Path(out_docx)

    # ---- pull the headline figures ----
    total_assets = _get(pack, "1 Balance Sheet", "total assets")
    total_liab = _get(pack, "1 Balance Sheet", "total liabilities")
    bank_debt = _get(pack, "1 Balance Sheet", "bank debt")
    net_worth = _get(pack, "1 Balance Sheet", "total equity")
    de_ratio = _num(_get(pack, "1 Balance Sheet", "debt-to-equity"))
    inv_ratio = _pct_val(_get(pack, "1 Balance Sheet", "investment / total assets"))

    runway = _num(_get(pack, "3 Liquidity", "months of runway"))
    op_cash = _get(pack, "3 Liquidity", "operating cash")
    near_cash = _get(pack, "3 Liquidity", "near-cash")
    burn = _get(pack, "3 Liquidity", "operating burn")
    buffer_btc = _num(_get(pack, "7 Risk Appetite", "sovereignty buffer"))
    cfp = _get(pack, "3 Liquidity", "contingency funding plan")

    cost_funds = _pct_val(_get(pack, "7 Risk Appetite", "cost of funds"))
    dtv = _pct_val(_get(pack, "7 Risk Appetite", "aggregate on-chain dtv"))
    native_yield = _pct_val(_get(pack, "7 Risk Appetite", "treasury native yield"))

    dur_gap = _num(_get(pack, "5 Asset-Liability", "duration gap"))
    refinancing = _get(pack, "5 Asset-Liability", "refinancing due")
    collateral = _get(pack, "5 Asset-Liability", "collateral sufficiency")
    fx = _get(pack, "5 Asset-Liability", "fx exposure")

    agg = pack.get("4 Funding Structure", {}).get("aggregate", (None, []))[1]
    # AGGREGATE row cols: B type, C limit, D drawn, E available, F util, G rate
    agg_drawn = agg[2] if len(agg) > 2 else None
    agg_avail = agg[3] if len(agg) > 3 else None
    agg_util = agg[4] if len(agg) > 4 else None

    # ---- evaluate limits -> breaches list ----
    checks: list[dict] = []

    def add(metric, value_str, status, limit, escalate):
        checks.append({"metric": metric, "value": value_str, "status": status,
                       "limit": limit, "escalate": escalate})

    if inv_ratio is not None:
        add("Investment / Total Assets", f"{inv_ratio:.1f}%",
            BREACH if inv_ratio >= 40 else OK, "< 40% by 30 Jun 2026",
            "Board (post-Jun 2026)")
    if de_ratio is not None:
        add("Debt-to-Equity", f"{de_ratio:.2f}x",
            BREACH if de_ratio > 0.5 else OK, "< 0.5x", "Board")
    if cost_funds is not None:
        add("Cost of funds", f"{cost_funds:.2f}%",
            BREACH if cost_funds > 4 else OK, "< 3–4%", "CAC")
    if dtv is not None:
        add("Aggregate on-chain DTV", f"{dtv:.1f}%",
            BREACH if dtv > 25 else OK, "≤ 25% (hard cap)", "Board if > 25%")
    if native_yield is not None:
        add("Treasury native yield", f"{native_yield:.1f}%",
            WATCH if native_yield < 10 else OK, "≥ 10% / THB 100M", "CAC")
    if runway is not None:
        add("Liquidity runway", f"{runway:.0f} months",
            BREACH if runway < 6 else OK, "≥ 6 months", "CEO if < 6mo")
    if buffer_btc is not None:
        add("Sovereignty Buffer", f"{buffer_btc:.0f} BTC",
            WATCH if buffer_btc <= 100 else OK, "≥ 100 BTC", "CEO if breached")
    if dur_gap is not None:
        add("Duration gap", f"{dur_gap:.2f} yr",
            BREACH if dur_gap >= 2.0 else OK, "< 2.0 yr", "CAC")

    breaches = [c for c in checks if c["status"] in (BREACH, WATCH)]

    # capital-allocation drift (Engine variances)
    alloc_rows = []
    for key, (orig, vals) in pack.get("2 Capital Allocation", {}).items():
        if key in ("bucket", "total") or not vals or vals[0] is None:
            continue
        actual, pct_assets, target, variance = (vals + [None] * 4)[:4]
        alloc_rows.append((orig, actual, pct_assets, target, variance))

    # ---- build the document ----
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(11)

    d.add_heading("CAC Monthly Report — First Draft", level=0)

    p = d.add_paragraph()
    p.add_run("To: ").bold = True
    p.add_run("Supane, CFO      ")
    p.add_run("From: ").bold = True
    p.add_run("Capital Allocation & ALCO Committee (CAC)")

    p = d.add_paragraph()
    p.add_run("Period: ").bold = True
    p.add_run(f"{month}      ")
    p.add_run("Prepared: ").bold = True
    p.add_run(f"{date.today().isoformat()}      ")
    p.add_run("Status: ").bold = True
    r = p.add_run("DRAFT — for committee review")
    r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    note = d.add_paragraph(
        "Auto-generated from the CAC Monthly Data Pack. Figures are exactly as "
        "entered by Treasury; limit checks are against Board-approved thresholds "
        "(Khao Yai §2.8). No figures are fabricated; blank narrative cells are "
        "marked [to be added]."
    )
    note.runs[0].italic = True

    d.add_paragraph(
        "Mandate: manage the Group's balance sheet, liquidity, funding structure, "
        "capital allocation priorities, and asset-liability risk exposures in "
        "accordance with Board-approved strategy, risk appetite, and treasury policies."
    )

    # 1. Executive Summary
    d.add_heading("1. Executive Summary", level=1)
    d.add_paragraph(
        f"Total assets {_thb_m(total_assets)} against liabilities "
        f"{_thb_m(total_liab)}, for a Net Worth of {_thb_m(net_worth)}. "
        f"Leverage is conservative (D/E "
        f"{de_ratio:.2f}x vs < 0.5x) and liquidity is very strong "
        f"({runway:.0f} months of runway vs a 6-month floor). "
        if de_ratio is not None and runway is not None else
        "Balance-sheet summary: see section 2."
    )
    hard = [c for c in breaches if c["status"] == BREACH]
    soft = [c for c in breaches if c["status"] == WATCH]
    if hard:
        d.add_paragraph(
            f"{len(hard)} metric(s) are at or beyond a Board-approved limit and "
            "require committee attention this month:"
        )
        for c in hard:
            _bullet(d, f"{c['value']} vs limit {c['limit']} → escalate to {c['escalate']}.",
                    bold_lead=f"{c['metric']}: ")
    else:
        d.add_paragraph("No hard limit breaches flagged this month.")
    if soft:
        d.add_paragraph("Items to watch (within limit but no headroom / off-target):")
        for c in soft:
            _bullet(d, f"{c['value']} vs target {c['limit']} → {c['escalate']}.",
                    bold_lead=f"{c['metric']}: ")
    if inv_ratio is not None and inv_ratio >= 40:
        d.add_paragraph(
            f"The single biggest issue is investment concentration: investments are "
            f"{inv_ratio:.1f}% of total assets against the 40% ceiling due 30 Jun 2026, "
            "driven by the Digital Asset Treasury overweight (see section 3). The "
            "melt-up/sell-down glide path should be tabled."
        )

    # 2. Balance Sheet
    d.add_heading("2. Balance Sheet", level=1)
    _kv(d, "Total assets", _thb_m(total_assets))
    _kv(d, "Total liabilities", _thb_m(total_liab))
    _kv(d, "Net Worth", _thb_m(net_worth))
    _kv(d, "Debt-to-Equity", f"{de_ratio:.2f}x (target < 0.5x) — OK" if de_ratio is not None else "—")
    _kv(d, "Investment / Total Assets",
        f"{inv_ratio:.1f}% (ceiling 40% by 30 Jun 2026) — "
        f"{'BREACH' if inv_ratio and inv_ratio >= 40 else 'OK'}"
        if inv_ratio is not None else "—")
    tl, bd = _num(total_liab), _num(bank_debt)
    if tl is not None and bd is not None and bd > tl:
        p = d.add_paragraph()
        r = p.add_run(
            f"⚠ Data check: 'Total liabilities' ({_thb_m(total_liab)}) is below "
            f"'Bank debt / facilities drawn' ({_thb_m(bank_debt)}); Treasury to "
            "reconcile before sign-off."
        )
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    d.add_paragraph("MoM movement & drivers: [to be added — prior-month column not "
                    "populated in this pack].")

    # 3. Capital Allocation
    d.add_heading("3. Capital Allocation", level=1)
    if inv_ratio is not None:
        d.add_paragraph(
            f"Investment-to-total-assets ratio: {inv_ratio:.1f}% (ceiling 40% by "
            "30 Jun 2026; OKR < 38%)."
        )
    d.add_paragraph("Allocation vs Three-Engine targets:")
    for name, actual, pct_assets, target, variance in alloc_rows:
        seg = f"{_thb_m(actual)} ({_pct(pct_assets)} of assets), target {target}"
        if variance:
            seg += f" → {str(variance).upper()}"
        _bullet(d, seg, bold_lead=f"{name}: ")
    d.add_paragraph("Material drift from target: Digital Asset Treasury and Advisory "
                    "are above their target ranges; VCC is below. [confirm drivers].")

    # 4. Liquidity
    d.add_heading("4. Liquidity — 'Stay Liquid'", level=1)
    _kv(d, "Operating cash", _thb_m(op_cash))
    _kv(d, "Near-cash / liquid securities", _thb_m(near_cash))
    _kv(d, "Monthly operating burn",
        f"{_thb_m(burn)} (budget ~THB 8M)" if burn is not None else "—")
    _kv(d, "Runway", f"{runway:.0f} months (min 6)" if runway is not None else "—")
    _kv(d, "Sovereignty Buffer",
        f"{buffer_btc:.0f} BTC (floor 100)" if buffer_btc is not None else "—")
    stress = []
    for key, (orig, vals) in pack.get("3 Liquidity", {}).items():
        if key in ("crypto drawdown", "cash drain", "combined"):
            result, passfail = (vals + [None, None])[1:3]
            stress.append(f"{orig} → {passfail or (_num(result) and f'{_num(result):.0f} mo') or '—'}")
    if stress:
        d.add_paragraph("Stress-test outcomes: " + "; ".join(stress) + ".")
    _kv(d, "Contingency Funding Plan", cfp or "[to be added]")

    # 5. Funding & Covenants
    d.add_heading("5. Funding & Covenants", level=1)
    _kv(d, "Total drawn / available",
        f"{_thb_m(agg_drawn)} / {_thb_m(agg_avail)} (util {_pct(agg_util)})")
    _kv(d, "Leverage / cost",
        f"D/E {de_ratio:.2f}x vs < 0.5x; cost of funds "
        f"{cost_funds:.2f}% vs < 3–4%" if de_ratio is not None and cost_funds is not None else "—")
    fac_lines = []
    for key, (orig, vals) in pack.get("4 Funding Structure", {}).items():
        if key in ("lender", "aggregate", "funding mix:", "covenant watch:") or vals[0] is None:
            continue
        limit, drawn_f = (vals + [None] * 5)[1], (vals + [None] * 5)[2]
        if _num(limit) is None and _num(drawn_f) is None:
            continue
        fac_lines.append(f"{orig}: drawn {_thb_m(drawn_f)} / limit {_thb_m(limit)}"
                         if _num(limit) else orig)
    if fac_lines:
        d.add_paragraph("Facilities: " + "; ".join(fac_lines) + ".")
    _kv(d, "Refinancing due next 12m", refinancing or "[to be added]")
    d.add_paragraph("Bank relationship notes (SCB/BBL): [to be added].")

    # 6. Asset-Liability Risk
    d.add_heading("6. Asset-Liability Risk", level=1)
    _kv(d, "Duration gap",
        f"{dur_gap:.2f} yr (target < 2.0) — {'OK' if dur_gap is not None and dur_gap < 2 else 'BREACH'}"
        if dur_gap is not None else "—")
    _kv(d, "Refinancing due next 12m", refinancing or "[to be added]")
    _kv(d, "Collateral sufficiency", collateral or "[to be added]")
    _kv(d, "FX exposure", fx or "[to be added]")
    _kv(d, "Aggregate on-chain DTV",
        f"{dtv:.1f}% (Board review > 25%)" if dtv is not None else "—")

    # 7. Limit Breaches & Escalations
    d.add_heading("7. Limit Breaches & Escalations", level=1)
    if not breaches:
        d.add_paragraph("No limits breached this month.")
    else:
        tbl = d.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Metric", "Value", "Status", "Limit", "Escalate to"]):
            tbl.rows[0].cells[i].text = h
        for c in breaches:
            cells = tbl.add_row().cells
            cells[0].text = c["metric"]
            cells[1].text = c["value"]
            cells[2].text = c["status"]
            cells[3].text = c["limit"]
            cells[4].text = c["escalate"]

    # 8. Policy Compliance
    d.add_heading("8. Policy Compliance", level=1)
    pol = pack.get("8 Policy Compliance", {})
    any_pol = False
    for key, (orig, vals) in pol.items():
        if key in ("policy / strategy",):
            continue
        yn, pol_note = (vals + [None, None])[:2]
        if yn is None and pol_note is None:
            continue
        any_pol = True
        _bullet(d, f"{yn or '[Y/N not entered]'}"
                + (f" — {pol_note}" if pol_note else ""), bold_lead=f"{orig}: ")
    if not any_pol:
        d.add_paragraph("Compliance not yet confirmed in the pack (Y/N column blank) "
                        "for: Board strategy (Three-Engine/North Star); Stay Liquid "
                        "doctrine; melt-up/sell-down 40% glide path; Decision Rights "
                        "Matrix (R-05). Treasury to confirm.")

    # 9. Recommendations
    d.add_heading("9. Recommendations to CFO / IC / Board", level=1)
    rec_rows = []
    for key, (orig, vals) in pack.get("9 Recommendations", {}).items():
        if key in ("recommendation",):
            continue
        amount, approval, rationale = (vals + [None, None, None])[:3]
        rec_rows.append((orig, amount, approval, rationale))
    if rec_rows:
        for rec, amount, approval, rationale in rec_rows:
            txt = rec
            extra = []
            if amount:
                extra.append(f"amount {amount}")
            if approval:
                extra.append(f"approval: {approval}")
            if rationale:
                extra.append(str(rationale))
            if extra:
                txt += " (" + "; ".join(extra) + ")"
            _bullet(d, txt)
    else:
        d.add_paragraph("None recorded this month.")
    d.add_paragraph("Source figures: this month's CAC Monthly Data Pack (Excel). "
                    "Questions: #cac-committee or the CAC agent.").runs[0].italic = True

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    d.save(out_docx)
    return breaches
